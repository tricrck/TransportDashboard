"""
Enhanced Data Fetcher Service
Supports all data source types with format detection and schema inference
"""

import requests
import pandas as pd
import json
import xml.etree.ElementTree as ET
import pyarrow.parquet as pq
from datetime import datetime
from flask import current_app
from models import DataSource, DataSourceType, AuthType, DataFormat, DataRefreshLog, db
import traceback
import io
import tempfile
import os


class DataFetcher:
    """
    Enhanced service for fetching data from all source types
    """
    
    @staticmethod
    def fetch_data(data_source, force_refresh=False):
        """
        Fetch data from a data source
        
        Args:
            data_source: DataSource object
            force_refresh: Force refresh even if cached
            
        Returns:
            dict: {success, data, from_cache, response_time, record_count}
        """
        try:
            # Check cache first
            if not force_refresh and data_source.is_cache_valid:
                return {
                    'success': True,
                    'data': data_source.cached_data,
                    'from_cache': True,
                    'cached_at': data_source.cached_at
                }
            
            # Start refresh log
            refresh_log = DataRefreshLog.start_refresh(
                data_source=data_source,
                triggered_by='manual' if force_refresh else 'auto'
            )
            
            start_time = datetime.utcnow()
            
            # Fetch based on source type
            if data_source.source_type == DataSourceType.API:
                result = DataFetcher._fetch_from_api(data_source)
            elif data_source.source_type == DataSourceType.UPLOAD:
                result = DataFetcher._fetch_from_upload(data_source)
            elif data_source.source_type == DataSourceType.LINK:
                result = DataFetcher._fetch_from_link(data_source)
            elif data_source.source_type == DataSourceType.DATABASE:
                result = DataFetcher._fetch_from_database(data_source)
            elif data_source.source_type == DataSourceType.DOCUMENT:
                result = DataFetcher._fetch_from_document(data_source)
            elif data_source.source_type == DataSourceType.SPREADSHEET:
                result = DataFetcher._fetch_from_spreadsheet(data_source)
            else:
                raise ValueError(f'Unsupported source type: {data_source.source_type}')
            
            # Calculate performance metrics
            end_time = datetime.utcnow()
            response_time = (end_time - start_time).total_seconds() * 1000  # ms
            
            # Process data
            processed_data = DataFetcher._process_data(result['data'], data_source)
            data_source.cache_data(processed_data)
            
            # Record success
            record_count = len(processed_data) if isinstance(processed_data, list) else 1
            data_source.record_success(
                response_time=response_time,
                record_count=record_count
            )
            
            if refresh_log:
                refresh_log.complete_success(
                    records_fetched=record_count,
                    records_processed=record_count,
                    data_size_bytes=len(str(processed_data))
                )
            
            return {
                'success': True,
                'data': processed_data,
                'from_cache': False,
                'response_time': response_time,
                'record_count': record_count
            }
            
        except Exception as e:
            error_msg = str(e)
            error_trace = traceback.format_exc()
            
            current_app.logger.error(f'Data fetch error for {data_source.name}: {error_msg}')
            current_app.logger.error(error_trace)
            
            # Record error
            data_source.record_error(error_msg)
            if 'refresh_log' in locals() and refresh_log:
                refresh_log.complete_error(error_msg, error_trace)
            
            return {
                'success': False,
                'error': error_msg,
                'data': None
            }
    
    @staticmethod
    def _fetch_from_api(data_source):
        """Fetch data from API endpoint"""
        try:
            headers = data_source.api_headers or {}
            params = data_source.api_params or {}
            
            # Add authentication
            auth = None
            if data_source.auth_type == AuthType.BASIC:
                auth = (data_source.auth_username, data_source.auth_password)
            elif data_source.auth_type == AuthType.BEARER:
                headers['Authorization'] = f'Bearer {data_source.auth_token}'
            elif data_source.auth_type == AuthType.API_KEY:
                headers['X-API-Key'] = data_source.auth_api_key
            elif data_source.auth_type == AuthType.QUERY_PARAM:
                params['api_key'] = data_source.auth_api_key
            
            # Make request
            timeout = data_source.api_timeout or 30
            
            if data_source.api_method == 'GET':
                response = requests.get(
                    data_source.api_endpoint,
                    headers=headers,
                    params=params,
                    auth=auth,
                    timeout=timeout
                )
            elif data_source.api_method in ['POST', 'PUT', 'PATCH']:
                response = requests.request(
                    data_source.api_method,
                    data_source.api_endpoint,
                    headers=headers,
                    params=params,
                    data=data_source.api_body,
                    auth=auth,
                    timeout=timeout
                )
            else:
                raise ValueError(f'Unsupported HTTP method: {data_source.api_method}')
            
            response.raise_for_status()
            
            # Parse response
            data = DataFetcher._parse_response(response, data_source.data_format)
            
            # Extract data path if specified
            if data_source.data_path:
                data = DataFetcher._extract_data_path(data, data_source.data_path)
            
            return {'success': True, 'data': data}
            
        except Exception as e:
            raise Exception(f'API request failed: {str(e)}')
    
    @staticmethod
    def _fetch_from_upload(data_source):
        """Fetch data from uploaded file"""
        if not data_source.file_path or not os.path.exists(data_source.file_path):
            raise ValueError('File not found')
        
        return DataFetcher._read_file(data_source.file_path, data_source.data_format)
    
    @staticmethod
    def _fetch_from_link(data_source):
        """Fetch data from external file URL"""
        try:
            headers = {}
            auth = None
            
            # Add authentication if required
            if data_source.auth_type == AuthType.BASIC:
                auth = (data_source.auth_username, data_source.auth_password)
            elif data_source.auth_type == AuthType.BEARER:
                headers['Authorization'] = f'Bearer {data_source.auth_token}'
            elif data_source.auth_type == AuthType.API_KEY:
                headers['X-API-Key'] = data_source.auth_api_key
            
            # Download file
            response = requests.get(
                data_source.file_url,
                headers=headers,
                auth=auth,
                timeout=60,
                stream=True
            )
            response.raise_for_status()
            
            # Save to temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{data_source.data_format.value}') as tmp:
                for chunk in response.iter_content(chunk_size=8192):
                    tmp.write(chunk)
                tmp_path = tmp.name
            
            try:
                # Read the temp file
                result = DataFetcher._read_file(tmp_path, data_source.data_format)
                return result
            finally:
                # Cleanup temp file
                if os.path.exists(tmp_path):
                    os.unlink(tmp_path)
                    
        except Exception as e:
            raise Exception(f'Link fetch failed: {str(e)}')
    
    @staticmethod
    def _fetch_from_database(data_source):
        """Fetch data from database"""
        try:
            import sqlalchemy
            from sqlalchemy import create_engine, text
            
            # Create engine
            engine = create_engine(data_source.db_connection_string)
            
            # Execute query
            query = data_source.query_string or f"SELECT * FROM {data_source.db_schema}.{data_source.db_table}"
            
            with engine.connect() as conn:
                result = conn.execute(text(query))
                
                # Convert to list of dicts
                columns = result.keys()
                data = [dict(zip(columns, row)) for row in result.fetchall()]
            
            return {'success': True, 'data': data}
            
        except Exception as e:
            raise Exception(f'Database query failed: {str(e)}')
    
    @staticmethod
    def _fetch_from_document(data_source):
        """Fetch data from document (PDF, DOCX, TXT)"""
        if not data_source.file_path:
            raise ValueError('No file path specified')
        
        try:
            if data_source.data_format == DataFormat.PDF:
                # Extract text from PDF
                import PyPDF2
                with open(data_source.file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() + "\n"
                return {'success': True, 'data': {'text': text, 'pages': len(reader.pages)}}
            
            elif data_source.data_format == DataFormat.DOCX:
                # Extract text from DOCX
                import docx
                doc = docx.Document(data_source.file_path)
                text = "\n".join([para.text for para in doc.paragraphs])
                return {'success': True, 'data': {'text': text, 'paragraphs': len(doc.paragraphs)}}
            
            elif data_source.data_format == DataFormat.TXT:
                with open(data_source.file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                return {'success': True, 'data': {'text': text}}
            
            else:
                raise ValueError(f'Unsupported document format: {data_source.data_format}')
                
        except Exception as e:
            raise Exception(f'Document reading failed: {str(e)}')
    
    @staticmethod
    def _fetch_from_spreadsheet(data_source):
        """Fetch data from spreadsheet"""
        return DataFetcher._fetch_from_upload(data_source)
    
    @staticmethod
    def _read_file(file_path, data_format):
        """Read file based on format"""
        try:
            if data_format == DataFormat.JSON:
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                return {'success': True, 'data': data}
            
            elif data_format == DataFormat.CSV:
                df = pd.read_csv(file_path)
                # Handle NaN values
                df = df.where(pd.notnull(df), None)
                data = df.to_dict('records')
                return {'success': True, 'data': data}
            
            elif data_format == DataFormat.XML:
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = DataFetcher._parse_xml(f.read())
                return {'success': True, 'data': data}
            
            elif data_format == DataFormat.EXCEL:
                df = pd.read_excel(file_path, engine='openpyxl')
                df = df.where(pd.notnull(df), None)
                data = df.to_dict('records')
                return {'success': True, 'data': data}
            
            elif data_format == DataFormat.PARQUET:
                table = pq.read_table(file_path)
                df = table.to_pandas()
                df = df.where(pd.notnull(df), None)
                data = df.to_dict('records')
                return {'success': True, 'data': data}
            
            elif data_format == DataFormat.HTML:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                # Try to extract tables
                tables = pd.read_html(io.StringIO(content))
                if tables:
                    data = tables[0].to_dict('records')
                else:
                    data = {'html': content}
                return {'success': True, 'data': data}
            
            else:
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = f.read()
                return {'success': True, 'data': data}
                
        except Exception as e:
            raise Exception(f'Error reading file: {str(e)}')
    
    @staticmethod
    def _parse_response(response, data_format):
        """Parse HTTP response based on format"""
        if data_format == DataFormat.JSON:
            return response.json()
        
        elif data_format == DataFormat.CSV:
            df = pd.read_csv(io.StringIO(response.text))
            df = df.where(pd.notnull(df), None)
            return df.to_dict('records')
        
        elif data_format == DataFormat.XML:
            return DataFetcher._parse_xml(response.text)
        
        elif data_format == DataFormat.HTML:
            try:
                tables = pd.read_html(response.text)
                if tables:
                    return tables[0].to_dict('records')
            except:
                pass
            return {'html': response.text}
        
        else:
            return response.text
    
    @staticmethod
    def _parse_xml(xml_string):
        """Parse XML to dictionary"""
        try:
            root = ET.fromstring(xml_string)
            return DataFetcher._xml_to_dict(root)
        except Exception as e:
            raise Exception(f'Error parsing XML: {str(e)}')
    
    @staticmethod
    def _xml_to_dict(element):
        """Convert XML element to dictionary"""
        result = {}
        
        # Add attributes
        if element.attrib:
            result.update(element.attrib)
        
        # Add text content
        if element.text and element.text.strip():
            if len(result) == 0:
                return element.text.strip()
            result['_text'] = element.text.strip()
        
        # Add child elements
        for child in element:
            child_dict = DataFetcher._xml_to_dict(child)
            if child.tag in result:
                if not isinstance(result[child.tag], list):
                    result[child.tag] = [result[child.tag]]
                result[child.tag].append(child_dict)
            else:
                result[child.tag] = child_dict
        
        return result
    
    @staticmethod
    def _extract_data_path(data, path):
        """Extract data from nested structure using JSONPath-like syntax"""
        try:
            # Simple implementation - supports dot notation
            parts = path.strip('$.').split('.')
            result = data
            
            for part in parts:
                if isinstance(result, dict):
                    result = result.get(part)
                elif isinstance(result, list) and part.isdigit():
                    result = result[int(part)]
                else:
                    raise ValueError(f'Cannot access path: {part}')
                
                if result is None:
                    break
            
            return result
        except Exception as e:
            raise Exception(f'Error extracting data path: {str(e)}')
    
    @staticmethod
    def _process_data(data, data_source):
        """Process and transform data"""
        # Apply transformations if specified
        if data_source.transform_script:
            try:
                # Safe execution environment (simplified)
                # In production, use RestrictedPython or similar
                exec_globals = {'data': data, 'pd': pd, 'json': json}
                exec(data_source.transform_script, exec_globals)
                data = exec_globals.get('result', data)
            except Exception as e:
                current_app.logger.error(f'Transform script error: {str(e)}')
        
        return data
    
    @staticmethod
    def test_connection(data_source):
        """
        Test data source connection
        
        Returns:
            dict: {success, message, response_time, record_count, data}
        """
        try:
            start_time = datetime.utcnow()
            result = DataFetcher.fetch_data(data_source, force_refresh=True)
            end_time = datetime.utcnow()
            
            response_time = (end_time - start_time).total_seconds() * 1000
            
            if result['success']:
                return {
                    'success': True,
                    'message': 'Connection successful',
                    'response_time': response_time,
                    'record_count': result.get('record_count', 0),
                    'data': result.get('data')  # Include data for schema inference
                }
            else:
                return {
                    'success': False,
                    'message': result.get('error', 'Connection failed'),
                    'response_time': response_time
                }
                
        except Exception as e:
            return {
                'success': False,
                'message': str(e),
                'response_time': None
            }